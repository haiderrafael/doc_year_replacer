import os
import utils
from docx import Document
import fitz

def test():
    # Test DOCX
    utils.replace_year_in_docx('test_doc.docx', 'updated_test_doc.docx', '2024', '2026')
    doc = Document('updated_test_doc.docx')
    text = '\n'.join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += '\n' + '\n'.join([p.text for p in cell.paragraphs])
                
    if '2026' in text and '2024' not in text:
        print("DOCX: Success")
    else:
        print("DOCX: Failed")
        print("Found text:", text)

    # Test PDF
    utils.replace_year_in_pdf('test_doc.pdf', 'updated_test_doc.pdf', '2024', '2026')
    pdf = fitz.open('updated_test_doc.pdf')
    pdf_text = ''
    for page in pdf:
        pdf_text += page.get_text()
        
    if '2026' in pdf_text and '2024' not in pdf_text:
        print("PDF: Success")
    else:
        print("PDF: Failed")
        print("Found PDF text:", pdf_text)

if __name__ == '__main__':
    test()
