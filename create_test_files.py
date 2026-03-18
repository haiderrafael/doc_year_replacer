import os
try:
    from docx import Document
    import fitz
except ImportError:
    print("Dependencies not installed yet.")
    exit(1)

# Create docx
doc = Document()
doc.add_heading('Documento de Prueba', 0)
doc.add_paragraph('Este es un documento del año 2024 que necesita ser actualizado.')
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Dato'
hdr_cells[1].text = 'Año: 2024'
doc.save('test_doc.docx')
print("Created test_doc.docx")

# Create pdf
pdf = fitz.open()
page = pdf.new_page()
page.insert_text(fitz.Point(72, 72), "Este es un reporte del 2024.", fontsize=14, fontname="helv")
pdf.save('test_doc.pdf')
print("Created test_doc.pdf")
