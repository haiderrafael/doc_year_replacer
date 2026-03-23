import os
from docx import Document
import fitz  # PyMuPDF

def replace_year_in_docx(input_filepath, output_filepath, old_year, new_year):
    """
    Replaces all occurrences of old_year with new_year in a docx document.
    Preserves styling accurately even when the year is split across multiple runs.
    """
    doc = Document(input_filepath)

    def _replace_in_paragraphs(paragraphs):
        for paragraph in paragraphs:
            if old_year not in paragraph.text:
                continue

            # Case 1: The year is fully contained within individual runs
            for run in paragraph.runs:
                if old_year in run.text:
                    run.text = run.text.replace(old_year, new_year)
            
            # Case 2: The year spans across multiple runs
            while old_year in paragraph.text:
                runs = paragraph.runs
                full_text = "".join(r.text for r in runs)
                start_idx = full_text.find(old_year)
                if start_idx == -1: 
                    break # Safety check
                
                # Find the run containing the start of the year
                curr_len = 0
                start_run_idx = -1
                char_in_start_run = -1
                for i, r in enumerate(runs):
                    r_len = len(r.text)
                    if curr_len <= start_idx < curr_len + r_len:
                        start_run_idx = i
                        char_in_start_run = start_idx - curr_len
                        break
                    curr_len += r_len
                    
                # Find the run containing the end of the year
                end_idx = start_idx + len(old_year)
                curr_len = 0
                end_run_idx = -1
                char_in_end_run = -1
                for i, r in enumerate(runs):
                    r_len = len(r.text)
                    if curr_len < end_idx <= curr_len + r_len:
                        end_run_idx = i
                        char_in_end_run = end_idx - curr_len
                        break
                    curr_len += r_len
                    
                if start_run_idx == -1 or end_run_idx == -1:
                    break # Should not happen, but avoid infinite loops
                    
                # Replace logic
                start_run = runs[start_run_idx]
                if start_run_idx == end_run_idx:
                    start_run.text = start_run.text[:char_in_start_run] + new_text + start_run.text[char_in_end_run:]
                else:
                    end_run = runs[end_run_idx]
                    start_run.text = start_run.text[:char_in_start_run] + new_year
                    for i in range(start_run_idx + 1, end_run_idx):
                        runs[i].text = ""
                    end_run.text = end_run.text[char_in_end_run:]

    _replace_in_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs)

    doc.save(output_filepath)

def replace_year_in_pdf(input_filepath, output_filepath, old_year, new_year):
    """
    Replaces occurrences of old_year with new_year in a PDF file using PyMuPDF.
    Redacts text and inserts the new text trying to match original font, size, and color.
    """
    doc = fitz.open(input_filepath)

    for page in doc:
        # Get detailed text dictionary to extract font size, name, and color
        text_dict = page.get_text("dict")
        font_info_map = {}
        
        for block in text_dict.get("blocks", []):
            if "lines" not in block: continue
            for line in block["lines"]:
                for span in line["spans"]:
                    color_int = span.get("color", 0)
                    # PyMuPDF color int is 0xRRGGBB
                    r = ((color_int >> 16) & 255) / 255.0
                    g = ((color_int >> 8) & 255) / 255.0
                    b = (color_int & 255) / 255.0
                    
                    bbox_tuple = tuple(span["bbox"])
                    font_info_map[bbox_tuple] = (
                        span.get("font", "helv"), 
                        span.get("size", 11), 
                        (r, g, b)
                    )

        # Search for instances
        text_instances = page.search_for(old_year)
        
        for inst in text_instances:
            # We must clear the existing text using redaction.
            page.add_redact_annot(inst, fill=(1, 1, 1)) # standard redaction
            
        page.apply_redactions()

        # Insert new text where the old instances were
        for inst in text_instances:
            rect = fitz.Rect(inst)
            
            # Find the best intersecting span for exact font properties
            fname, fsize, color = "helv", 11, (0, 0, 0)
            found_match = False
            best_overlap = -1
            
            for bbox, info in font_info_map.items():
                span_rect = fitz.Rect(bbox)
                if span_rect.intersects(rect):
                    overlap_rect = span_rect.intersect(rect)
                    overlap_area = overlap_rect.get_area()
                    if overlap_area > best_overlap:
                        best_overlap = overlap_area
                        fname, fsize, color = info
                        found_match = True
                    
            if not found_match:
                fsize = (rect.y1 - rect.y0) * 0.85
                if fsize < 8: fsize = 11
                
            # Broadly map specific embedded fonts to base-14 PyMuPDF insertion fonts ensuring style preservation
            fname_lower = fname.lower()
            is_bold = "bold" in fname_lower or "black" in fname_lower or "heavy" in fname_lower
            is_italic = "italic" in fname_lower or "oblique" in fname_lower
            
            insert_font = "helv"
            if "times" in fname_lower or "serif" in fname_lower:
                if is_bold and is_italic: insert_font = "tibi"
                elif is_bold: insert_font = "tibo"
                elif is_italic: insert_font = "tiit"
                else: insert_font = "tiro"
            elif "cour" in fname_lower or "mono" in fname_lower:
                if is_bold and is_italic: insert_font = "cobo"
                elif is_bold: insert_font = "cob"
                elif is_italic: insert_font = "coob"
                else: insert_font = "cour"
            else: # Default sans-serif fallback (Arial, Calibri, Verdana, Helvetica...)
                if is_bold and is_italic: insert_font = "hebi"
                elif is_bold: insert_font = "hebo"
                elif is_italic: insert_font = "heit"
                else: insert_font = "helv"

            p = fitz.Point(rect.x0, rect.y1 - (rect.height * 0.2))
            
            try:
                page.insert_text(p, new_year, fontsize=fsize, fontname=insert_font, color=color)
            except Exception:
                page.insert_text(p, new_year, fontsize=fsize, fontname="helv", color=color)
            
    doc.save(output_filepath)
